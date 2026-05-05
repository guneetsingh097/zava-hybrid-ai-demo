"""Generate the Zava Insurance Hybrid AI Demo talk track as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title
title = doc.add_heading("Zava Insurance: Hybrid AI Demo", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("On-Device AI + Microsoft WorkIQ — Best of Both Worlds on Windows 11")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x00, 0x47, 0x8A)

doc.add_paragraph()

# Meta info
meta = doc.add_table(rows=5, cols=2)
meta.style = 'Light Grid Accent 1'
meta_data = [
    ("Duration", "2–3 minutes"),
    ("Audience", "Enterprise IT Decision Makers (ITDM) & Business Decision Makers (BDM)"),
    ("Key Message", "Windows 11 Copilot+ PCs deliver hybrid AI — fast private inference on the NPU + secure organizational intelligence from Microsoft 365 via WorkIQ"),
    ("Hardware", "Any Copilot+ PC (Snapdragon X, Intel Core Ultra, AMD Ryzen AI)"),
    ("Prerequisites", "Windows 11, phi-npu.exe (Phi Silica), Python 3.11+, Foundry Local"),
]
for i, (k, v) in enumerate(meta_data):
    meta.rows[i].cells[0].text = k
    meta.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_heading("The Story", level=1)
doc.add_paragraph(
    "An insurance company deploys AI-powered claims processing to 10,000 field adjusters. "
    "The challenge: they need AI that works anywhere (in the field, offline) while still "
    "connecting securely to organizational knowledge when available. The answer is hybrid AI "
    "on Windows 11 — the NPU handles all sensitive inference locally (zero cloud cost, zero "
    "data exposure), and Microsoft WorkIQ enriches responses with CRM data, claims history, "
    "and internal communications when connected."
)
doc.add_paragraph(
    "This is only possible on Windows. The dedicated NPU hardware means AI runs at hardware "
    "speed without touching the CPU or GPU. And the secure connection to Microsoft 365 means "
    "organizational intelligence flows in without sensitive customer data flowing out."
)

doc.add_heading("Demo Flow (2–3 minutes)", level=1)

# STEP 1
doc.add_heading("Step 1: Set the Stage (15 seconds)", level=2)
doc.add_paragraph("TALK TRACK:", style='Intense Quote')
doc.add_paragraph(
    '"Let me show you how a Windows 11 Copilot+ PC can run AI-powered insurance claims '
    'processing — entirely on the device, with zero cloud cost. This is Zava Insurance\'s '
    'hybrid AI prototype. Notice the NPU status indicator in the top right — green means '
    'the Neural Processing Unit is active and ready."'
)
doc.add_paragraph("ACTION: Show the Home tab. Point out the NPU Online indicator and WorkIQ Off status.")
# Insert screenshot
img_path = os.path.join(os.path.dirname(__file__), "demo_assets", "01_app_overview.png")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# STEP 2
doc.add_heading("Step 2: Claims AI — Photo Assessment on NPU (45 seconds)", level=2)
doc.add_paragraph("TALK TRACK:", style='Intense Quote')
doc.add_paragraph(
    '"A field adjuster photographs water damage at a customer\'s home. Watch what happens — '
    'I upload the photo and the NPU analyzes it instantly. No internet needed, no data leaves '
    'the device. The AI identifies damage type, estimates severity, and recommends next steps."'
)
doc.add_paragraph(
    "ACTION:\n"
    "1. Click 'Claims AI' tab\n"
    "2. Drag a water damage photo from sample_data/claim_photos/ into the upload area\n"
    "3. Select 'Water Damage' as damage type\n"
    "4. Click 'Assess with On-Device AI'\n"
    "5. Point out the response time (~8-12 seconds) and that it says 'Engine: NPU'\n"
    "6. Open Task Manager (Ctrl+Shift+Esc) → Performance tab → show NPU utilization spike"
)
if os.path.exists(os.path.join(os.path.dirname(__file__), "demo_assets", "02_claims_ai.png")):
    doc.add_picture(os.path.join(os.path.dirname(__file__), "demo_assets", "02_claims_ai.png"), width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

doc.add_paragraph("💡 KEY POINT: ", style='List Bullet')
p = doc.paragraphs[-1]
p.add_run("Open Task Manager (Ctrl+Shift+Esc) → Performance tab. You'll see the NPU utilization spike during inference. This is dedicated silicon — it's not touching the CPU or GPU, which remain free for other work.")
doc.add_paragraph()

# STEP 3
doc.add_heading("Step 3: Policy Assistant — Hybrid AI in Action (45 seconds)", level=2)
doc.add_paragraph("TALK TRACK:", style='Intense Quote')
doc.add_paragraph(
    '"Now here\'s where it gets interesting. The Policy Assistant can answer questions using '
    'just the NPU — fast and private. But watch what happens when I connect to Microsoft WorkIQ..."'
)
doc.add_paragraph(
    "ACTION:\n"
    "1. Click 'Policy Assistant' tab\n"
    "2. First, ask WITHOUT WorkIQ: Type 'What does the HO-46 endorsement cover?' → Show fast NPU response\n"
    "3. Now check the 'Connect securely to Microsoft WorkIQ' checkbox\n"
    "4. Notice the WorkIQ indicator in the top bar turns green\n"
    "5. Ask: 'What is the current claim status?' → Show enriched response with CRM data\n"
    "6. Point out the 'Source: Microsoft 365 (WorkIQ)' badge on the response"
)
doc.add_paragraph("TALK TRACK (continued):", style='Intense Quote')
doc.add_paragraph(
    '"See the difference? Same NPU, same on-device processing — but now it\'s enriched with '
    'real organizational context. The AI pulled claim status, adjuster notes, and timeline '
    'from Microsoft 365. The sensitive reasoning still happened locally on the NPU — '
    'WorkIQ just provided the context securely. Best of both worlds."'
)
if os.path.exists(os.path.join(os.path.dirname(__file__), "demo_assets", "03_policy_assistant.png")):
    doc.add_picture(os.path.join(os.path.dirname(__file__), "demo_assets", "03_policy_assistant.png"), width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# STEP 4
doc.add_heading("Step 4: NPU Dashboard — The Business Case (30 seconds)", level=2)
doc.add_paragraph("TALK TRACK:", style='Intense Quote')
doc.add_paragraph(
    '"Let me show you why this matters at scale. The NPU Dashboard tracks every inference — '
    'zero cloud cost, zero network latency, full data sovereignty. If you deploy this to '
    '10,000 field adjusters doing 50 AI inferences per day, that\'s 130 million inferences '
    'per year at zero incremental cloud cost. That\'s $1.3 million saved annually — and that\'s '
    'just one application."'
)
doc.add_paragraph(
    "ACTION:\n"
    "1. Click 'NPU Dashboard' tab\n"
    "2. Point out: Cloud Cost Saved, Total Inferences, Avg Latency\n"
    "3. Show the Fleet-Scale Projection calculator (already shows $1.3M)\n"
    "4. Mention the Carbon Savings Calculator below"
)
if os.path.exists(os.path.join(os.path.dirname(__file__), "demo_assets", "05_npu_dashboard.png")):
    doc.add_picture(os.path.join(os.path.dirname(__file__), "demo_assets", "05_npu_dashboard.png"), width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# STEP 5
doc.add_heading("Step 5: Developer Experience — Customize with GitHub Copilot CLI (30 seconds)", level=2)
doc.add_paragraph("TALK TRACK:", style='Intense Quote')
doc.add_paragraph(
    '"One more thing. Because this is a Windows application built with standard web technologies, '
    'anyone in your organization can customize it using natural language. Let me show you — '
    'I\'ll use GitHub Copilot CLI to add a dark mode toggle to the app in seconds."'
)
doc.add_paragraph(
    "ACTION:\n"
    "1. Open a terminal next to the app\n"
    "2. Type: copilot \"Add a light mode / dark mode toggle button to the top menu bar "
    "next to the NPU status indicator. When toggled, switch all backgrounds, text, and cards "
    "to a dark theme.\"\n"
    "3. Show Copilot CLI making the code changes in real-time\n"
    "4. Refresh the browser — show the new toggle in the top menu\n"
    "5. Click it — show the app switch to dark mode"
)
doc.add_paragraph("TALK TRACK (closing):", style='Intense Quote')
doc.add_paragraph(
    '"That\'s the power of building on Windows. Dedicated AI hardware that runs offline, '
    'secure connection to organizational intelligence via Microsoft 365, and any developer '
    'in your company can extend it using natural language with GitHub Copilot. '
    'This is hybrid AI — only on Windows 11."'
)

doc.add_paragraph()
doc.add_heading("Key Messages for ITDM/BDM Audience", level=1)

bullets = [
    "Security & Compliance: All AI inference runs on-device. Sensitive customer data (PII, claims, photos) never leaves the hardware. Full data sovereignty.",
    "Cost at Scale: Zero incremental cloud cost per inference. At fleet scale (10K+ devices), this translates to millions in annual savings vs. cloud-only AI.",
    "Works Offline: Field adjusters can use AI in basements, rural areas, disaster zones — anywhere without connectivity. The NPU doesn't need internet.",
    "Best of Both Worlds: When connected, Microsoft WorkIQ securely enriches responses with organizational context (emails, Teams, documents) — without sending customer data to the cloud.",
    "Hardware Investment: Copilot+ PCs with dedicated NPUs are a platform for hundreds of AI applications — insurance is just one example.",
    "Developer Velocity: Standard web technologies + GitHub Copilot CLI means any developer can build, customize, and extend AI applications using natural language.",
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()
doc.add_heading("Technical Architecture", level=1)
doc.add_paragraph(
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│  Windows 11 Copilot+ PC                                        │\n"
    "│  ┌─────────────────┐  ┌──────────────────┐  ┌──────────────┐  │\n"
    "│  │   NPU (Primary)  │  │  CPU (Fallback)   │  │  GPU (Free)  │  │\n"
    "│  │   Phi Silica     │  │  Foundry Local    │  │  User apps   │  │\n"
    "│  │   ~8-12s/query   │  │  Qwen 2.5 1.5B   │  │              │  │\n"
    "│  └────────┬─────────┘  └────────┬──────────┘  └──────────────┘  │\n"
    "│           │                      │                               │\n"
    "│           └──────────┬───────────┘                               │\n"
    "│                      │                                           │\n"
    "│              ┌───────┴────────┐                                  │\n"
    "│              │  Flask App     │◄── Microsoft WorkIQ (optional)   │\n"
    "│              │  localhost:5000│    (M365 CRM context)            │\n"
    "│              └────────────────┘                                  │\n"
    "└─────────────────────────────────────────────────────────────────┘"
)

doc.add_paragraph()
doc.add_heading("Setup Instructions (Other Machines)", level=1)
doc.add_paragraph("1. Clone the repository: git clone https://github.com/gusing_microsoft/zava-hybrid-ai-demo")
doc.add_paragraph("2. Install Python dependencies: pip install -r requirements.txt")
doc.add_paragraph("3. Ensure phi-npu.exe is available (ships with Windows AI APIs on Copilot+ PCs)")
doc.add_paragraph("4. Install Foundry Local and pull the model: foundry model pull qwen2.5-1.5b")
doc.add_paragraph("5. Copy sample_data/workiq/*.md files to your OneDrive for Business folder 'Zava Insurance Demo'")
doc.add_paragraph("6. Run the app: python app.py")
doc.add_paragraph("7. Open http://localhost:5000 in your browser")

doc.add_paragraph()
doc.add_heading("Task Manager Tip", level=1)
doc.add_paragraph(
    "During the demo, keep Task Manager open (Ctrl+Shift+Esc → Performance tab). "
    "When AI inference runs, you'll see the NPU utilization spike while CPU and GPU remain idle. "
    "This visually demonstrates that AI is running on dedicated silicon — not stealing resources "
    "from other applications. This is a powerful visual for technical audiences."
)

# Save
output_path = os.path.join(os.path.dirname(__file__), "Zava_Insurance_Hybrid_AI_Demo_Guide.docx")
doc.save(output_path)
print(f"✅ Demo guide saved to: {output_path}")
